from docx import Document

from src.fas.fas_text_analysis import TextSummary


def extract_docx_data(path):
    try:
        doc = Document(path)

        # This can be used when finding keywords for summary, analysis, and or ranking depending on implememntation
        text = "\n".join(p.text for p in doc.paragraphs)

        analyzer = TextSummary(text) if text.strip() else None

        # Metadata extraction
        core = doc.core_properties
        metadata = {
            "author": core.author,
            "title": core.title,
            "subject": core.subject,
            "created": core.created,
            "modified": core.modified,
            "keywords": core.keywords,
            "category": core.category,
            "comments": core.comments,
            # This is to get stats of the content of the document
            "num_paragraphs": len(doc.paragraphs),
            "num_tables": len(doc.tables),
            "num_chars": len(text),
            "num_words": len(text.split()),
        }

        # If text exists, add text analysis
        if analyzer:
            metadata.update(analyzer.generate_text_analysis_data(10, 3))

        return metadata
    except Exception as e:
        return {"error": str(e)}
